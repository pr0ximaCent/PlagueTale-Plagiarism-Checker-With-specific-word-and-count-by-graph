import docx

doc = docx.Document("test_doc_2.docx")
for paragraph in doc.paragraphs:
    print(paragraph.text)


import docx

doc = docx.Document("test_doc_2.docx")
for paragraph in doc.paragraphs:
    print(paragraph.text)



def text_to_dictionary(text):
    dictionary = {}
    for word in text.split():
        if word in dictionary:
            dictionary[word] += 1
        else:
            dictionary[word] = 1
    return dictionary


    data1=text_to_dictionary(first.text)
    data2=text_to_dictionary(second.text)
    common = {word: min(data1[word], data2[word]) for word in data1 if word in data2 and len(word) > 4}
    most_frequent = sorted(common.items(), key=lambda x: x[1], reverse=True)[:10]
    print([{word: freq} for word, freq in most_frequent])